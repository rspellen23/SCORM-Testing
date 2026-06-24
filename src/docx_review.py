"""Course IR -> branded SME review .docx  (Gate 3 deliverable).

Per COURSE_CREATION_SYSTEM.md §3: the canonical artifact is the `.md`; the `.docx`
is its *readable face* for SME review. This renders a clean, branded Word doc from
the SAME IR the build consumes, so the thing reviewed matches the thing built.

Contains (per the Gate-3 spec):
  - course/unit title + the preamble context (Subject / Length / Objectives)
  - each slide as a Word Heading 2 + body (paragraphs, bullets, tables)
  - knowledge checks in a shaded box with the correct option marked
Markdown is canonical; this projection is regenerated whenever the .md changes.

Reuses html_paras / plain / parse_table from pptx_export (one HTML parser, not two).
"""
import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pptx_export import html_paras, plain, parse_table

NAVY = RGBColor(0x00, 0x3E, 0x51)
GREY = RGBColor(0x33, 0x33, 0x33)
BODY_FONT = "Open Sans"


def _shade(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _runs(p, runs, size=11, color=GREY, bold=False, italic=False):
    for r in runs:
        run = p.add_run(r["text"])
        run.bold = bool(r["bold"] or bold)
        run.italic = bool(r["italic"] or italic)
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.name = "Consolas" if r["code"] else BODY_FONT


def _para(doc, fragment, size=11, color=GREY, bold=False, italic=False, style=None):
    paras = html_paras(fragment) or [[{"text": "", "bold": False, "italic": False, "code": False}]]
    out = None
    for runs in paras:
        p = doc.add_paragraph(style=style)
        _runs(p, runs, size, color, bold, italic)
        out = p
    return out


def _accent_rgb(brand):
    hexv = (brand.accent if brand else "#3B82F6").lstrip("#")
    return RGBColor.from_string(hexv), hexv


def _preamble(md_path, which):
    """Pull the per-microlearning author meta lines (Subject/Length/Objectives) plus the
    instructional-design rationale from the raw .md so reviewers get provenance and can see
    WHY the unit/curriculum is structured the way it is. Best-effort, never raises."""
    try:
        text = open(md_path, encoding="utf-8").read()
    except OSError:
        return []
    secs = re.split(r"^##\s+Microlearning\s+", text, flags=re.M)
    head = secs[0] if secs else ""              # file preamble (course-level meta)
    lines = []
    # course-level bold meta + the curriculum-level "why these units / this order" rationale
    for key in ("Subject", "Estimated Length", "Learning Objectives", "Confidence Score",
                "Curriculum Rationale"):
        m = re.search(rf"^\*\*{re.escape(key)}:\*\*\s*(.+)$", head, flags=re.M)
        if m:
            lines.append((key, m.group(1).strip()))
    # per-unit "why this unit is built this way" rationale, from its Build-Notes block (plain line)
    sec = secs[which] if 0 < which < len(secs) else ""
    rm = re.search(r"Design Rationale:\s*(.+?)(?:\n(?:[A-Za-z][\w /]*:|\*\*)|\Z)", sec, flags=re.S)
    if rm:
        val = re.sub(r"\s*\n\s*", "; ", rm.group(1).strip()).strip("; ").strip()
        if val:
            lines.append(("Design Rationale", val))
    return lines


def _kc_box(doc, b, accent_rgb, accent_hex):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _shade(cell, "F2F7F4")
    cell.paragraphs[0].text = ""
    lab = cell.paragraphs[0]
    r = lab.add_run("KNOWLEDGE CHECK")
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = accent_rgb; r.font.name = BODY_FONT
    q = cell.add_paragraph()
    _runs(q, (html_paras(b.get("prompt")) or [[]])[0], 11, NAVY, bold=True)
    for i, o in enumerate(b.get("options", [])):
        p = cell.add_paragraph()
        letter = chr(ord("A") + i)
        mark = "  ✓" if o.get("correct") else ""
        head = p.add_run(f"{letter}) ")
        head.font.size = Pt(11); head.font.name = BODY_FONT
        head.font.color.rgb = accent_rgb if o.get("correct") else GREY
        head.bold = bool(o.get("correct"))
        _runs(p, (html_paras(o.get("html")) or [[]])[0], 11,
              NAVY if o.get("correct") else GREY, bold=bool(o.get("correct")))
        if mark:
            mr = p.add_run(mark); mr.bold = True; mr.font.color.rgb = accent_rgb
            mr.font.size = Pt(11); mr.font.name = BODY_FONT
    if b.get("feedback"):
        fb = cell.add_paragraph()
        lr = fb.add_run("Feedback: ")
        lr.bold = True; lr.italic = True; lr.font.size = Pt(10); lr.font.name = BODY_FONT
        lr.font.color.rgb = GREY
        fr = fb.add_run(b["feedback"]); fr.italic = True; fr.font.size = Pt(10)
        fr.font.color.rgb = GREY; fr.font.name = BODY_FONT
    doc.add_paragraph()


def _fmt_num(v):
    if v is None:
        return ""
    if isinstance(v, float) and not v.is_integer():
        return f"{v:g}"
    return str(int(v)) if isinstance(v, (int, float)) else str(v)


def _render_chart_review(doc, b, accent_hex):
    """Surface a chart in the SME review doc as a LABEL + a data table + its source,
    so the reviewer can verify every figure (and catch a missing/wrong source) before
    anything is published."""
    kind = {"bar": "Bar", "groupedBar": "Grouped bar", "stackedBar": "Stacked bar",
            "line": "Line", "pie": "Pie"}.get(b.get("chart"), "Chart")
    lab = doc.add_paragraph()
    r = lab.add_run(f"[{kind} chart] "); r.bold = True; r.font.color.rgb = NAVY
    r.font.name = BODY_FONT; r.font.size = Pt(11)
    if b.get("title"):
        t = lab.add_run(plain(b["title"])); t.font.color.rgb = NAVY
        t.font.name = BODY_FONT; t.font.size = Pt(11); t.bold = True
    cats = b.get("categories") or []
    series = [s for s in (b.get("series") or []) if isinstance(s, dict)]
    if cats and series:
        tbl = doc.add_table(rows=len(cats) + 1, cols=len(series) + 1)
        tbl.style = "Table Grid"
        hdr = [""] + [s.get("name") or f"Series {i + 1}" for i, s in enumerate(series)]
        for ci, txt in enumerate(hdr):
            cell = tbl.cell(0, ci); p = cell.paragraphs[0]
            run = p.add_run(txt); run.bold = True; run.font.size = Pt(10)
            run.font.name = BODY_FONT; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade(cell, accent_hex)
        for ri, cat in enumerate(cats, 1):
            cells = [cat] + [_fmt_num((s.get("data") or [None] * len(cats))[ri - 1]
                                      if ri - 1 < len(s.get("data") or []) else None) for s in series]
            for ci, txt in enumerate(cells):
                cell = tbl.cell(ri, ci); p = cell.paragraphs[0]
                run = p.add_run(txt); run.font.size = Pt(10); run.font.name = BODY_FONT
                run.bold = (ci == 0); run.font.color.rgb = GREY
    # source line — flag it RED if missing, since the build would reject it
    sp = doc.add_paragraph()
    if (b.get("source") or "").strip():
        s1 = sp.add_run("Source: "); s1.bold = True
        s2 = sp.add_run(plain(b["source"])); s2.italic = True
        for rr in (s1, s2):
            rr.font.size = Pt(9.5); rr.font.name = BODY_FONT; rr.font.color.rgb = GREY
    else:
        warn = sp.add_run("⚠ NO SOURCE — verify these figures and add a source before publishing.")
        warn.bold = True; warn.font.size = Pt(9.5); warn.font.name = BODY_FONT
        warn.font.color.rgb = RGBColor(0xA3, 0x20, 0x19)
    doc.add_paragraph()


def _render_table(doc, fragment, accent_hex):
    rows = parse_table(fragment)
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncol)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(ncol):
            cell = tbl.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            txt = row[ci] if ci < len(row) else ""
            run = p.add_run(txt)
            run.font.size = Pt(10); run.font.name = BODY_FONT
            if ri == 0:
                run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _shade(cell, accent_hex)
            else:
                run.font.color.rgb = GREY
    doc.add_paragraph()


def render_review_docx(ir, out_path, brand=None, md_path=None, which=1):
    """Render one microlearning IR to a branded review .docx at out_path."""
    accent_rgb, accent_hex = _accent_rgb(brand)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = GREY

    # ---- title
    h = doc.add_paragraph()
    r = h.add_run(ir.get("title", "Course"))
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = accent_rgb; r.font.name = BODY_FONT

    sub = doc.add_paragraph()
    sr = sub.add_run("SME Review Draft — markdown is canonical; this Word doc is its readable face. "
                     "Comment or use tracked changes; edits are applied back to the .md.")
    sr.italic = True; sr.font.size = Pt(9); sr.font.color.rgb = GREY; sr.font.name = BODY_FONT

    # ---- context box (preamble meta)
    meta = _preamble(md_path, which) if md_path else []
    if meta:
        tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
        cell = tbl.cell(0, 0); _shade(cell, "EEF1F2"); cell.paragraphs[0].text = ""
        cap = cell.paragraphs[0]
        cr = cap.add_run("Reviewer context — not shown to learners")
        cr.bold = True; cr.font.size = Pt(9); cr.font.color.rgb = NAVY; cr.font.name = BODY_FONT
        for k, v in meta:
            p = cell.add_paragraph()
            kr = p.add_run(f"{k}: "); kr.bold = True; kr.font.size = Pt(10)
            kr.font.color.rgb = NAVY; kr.font.name = BODY_FONT
            vr = p.add_run(v); vr.font.size = Pt(10); vr.font.color.rgb = GREY; vr.font.name = BODY_FONT
    doc.add_paragraph()

    # ---- slides
    slide_n = 0
    for b in ir.get("blocks", []):
        t = b.get("type")
        if t == "heading" and b.get("level", 2) <= 2:
            slide_n += 1
            hp = doc.add_heading(level=2)
            run = hp.add_run(f"Slide {slide_n} — {plain(b.get('html'))}")
            run.font.color.rgb = NAVY; run.font.name = BODY_FONT
            continue
        if t in ("sectionStart", "sectionEnd", "divider", "transition", "continue"):
            continue
        if t == "heading":                                   # level-3 subheading
            _para(doc, b.get("html"), size=12, color=NAVY, bold=True)
        elif t in ("paragraph", "headingParagraph"):
            if t == "headingParagraph":
                _para(doc, b.get("headingHtml"), size=12, color=NAVY, bold=True)
            _para(doc, b.get("html"))
        elif t == "list":
            ordered = b.get("ordered")
            for item in b.get("items", []):
                _para(doc, item, style="List Number" if ordered else "List Bullet")
        elif t == "note":
            _para(doc, "Note: " + (b.get("html") or ""), size=10, color=GREY, italic=True)
        elif t == "statement":
            _para(doc, b.get("html"), size=13, color=accent_rgb, bold=True)
        elif t in ("image", "imageText"):
            alt = b.get("alt") or b.get("caption") or "image"
            cap = doc.add_paragraph()
            cr = cap.add_run(f"[Visual: {alt}]")
            cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = accent_rgb; cr.font.name = BODY_FONT
            if t == "imageText" and b.get("html"):
                _para(doc, b.get("html"))
        elif t == "table":
            _render_table(doc, b.get("html"), accent_hex)
        elif t == "chart":
            _render_chart_review(doc, b, accent_hex)
        elif t == "cardGrid":
            for c in b.get("cards", []):
                p = doc.add_paragraph(style="List Bullet")
                tr = p.add_run(c.get("title") or ""); tr.bold = True
                tr.font.color.rgb = NAVY; tr.font.name = BODY_FONT; tr.font.size = Pt(11)
                if c.get("teaser"):
                    zr = p.add_run(" — " + c["teaser"]); zr.font.size = Pt(11)
                    zr.font.color.rgb = GREY; zr.font.name = BODY_FONT
        elif t == "knowledgeCheck":
            _kc_box(doc, b, accent_rgb, accent_hex)

    doc.save(out_path)
    return {"slides": slide_n}
