#!/usr/bin/env python3
"""Extract plain text from every accepted source doc in a project folder and
print one combined corpus to stdout, with a header before each file so the
drafter can tell the sources apart.

Accepted: .docx (python-docx), .md/.markdown/.txt (read as-is).
Unsupported (.pdf and anything else): a warning to stderr, skipped — never
silently dropped. Exit non-zero if the folder yields zero accepted docs.

Usage: extract_text.py <project-folder>
"""
import os, sys

ACCEPTED_TEXT = {".md", ".markdown", ".txt"}


def docx_text(path):
    """Paragraphs + table cell text, in document order, via python-docx."""
    from docx import Document
    doc = Document(path)
    out = []
    for p in doc.paragraphs:
        if p.text.strip():
            out.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                out.append(" | ".join(cells))
    return "\n".join(out)


def main():
    if len(sys.argv) != 2:
        sys.stderr.write("usage: extract_text.py <project-folder>\n")
        return 2
    root = sys.argv[1]
    if not os.path.isdir(root):
        sys.stderr.write(f"not a directory: {root}\n")
        return 2

    files = []
    for dirpath, _dirs, names in os.walk(root):
        for n in sorted(names):
            if n.startswith("."):           # skip .processing, .DS_Store, etc.
                continue
            files.append(os.path.join(dirpath, n))
    files.sort()

    found = 0
    for path in files:
        ext = os.path.splitext(path)[1].lower()
        rel = os.path.relpath(path, root)
        try:
            if ext == ".docx":
                text = docx_text(path)
            elif ext in ACCEPTED_TEXT:
                with open(path, encoding="utf-8", errors="replace") as fh:
                    text = fh.read()
            else:
                sys.stderr.write(f"SKIP (unsupported type {ext or 'none'}): {rel}\n")
                continue
        except Exception as e:                # noqa: BLE001 — report and keep going
            sys.stderr.write(f"SKIP (read error: {e}): {rel}\n")
            continue

        if not text.strip():
            sys.stderr.write(f"SKIP (empty): {rel}\n")
            continue

        found += 1
        print(f"\n\n===== SOURCE FILE: {rel} =====\n")
        print(text.rstrip())

    if found == 0:
        sys.stderr.write("ERROR: no accepted source documents found.\n")
        return 1
    sys.stderr.write(f"extracted {found} source document(s)\n")
    return 0


if __name__ == "__main__":
    sys.exit(main())
